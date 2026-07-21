"""Build a Torah reading handout (docx) for a Chavurat Or'Leans Shabbat service.

For a given service date, pulls the triennial-cycle aliyot from Hebcal,
fetches Hebrew (with cantillation) and English (JPS 2006) text from Sefaria,
and lays them out verse-by-verse in the house handout format:

    Parashat <Name>            (centered, 14pt)
    1st Aliyah                 (centered, 12pt bold)
    <Hebrew verse>             (RTL, 16pt)
    <English verse>            (LTR, 12pt)
    ...

Usage:
    python build_handout.py --date 2026-10-24 [--date 2026-11-07 ...]
                            [--aliyot 3] [--out ./handouts]

Output filename: "YYYY.MM.DD - <Parsha>.docx" (matches the Sefaria Source
Sheets naming used since April 2026).

Requires: python-docx (pip install python-docx). Network access to
hebcal.com and sefaria.org.
"""

import argparse
import html
import json
import re
import sys
import urllib.parse
import urllib.request
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ENGLISH_VERSION = "The Contemporary Torah, Jewish Publication Society, 2006"
USER_AGENT = "chavuratorleans-handouts/1.0 (chavuratorleans@gmail.com)"
ORDINALS = {1: "1st", 2: "2nd", 3: "3rd", 4: "4th", 5: "5th", 6: "6th", 7: "7th"}


def fetch_json(url):
    req = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
    with urllib.request.urlopen(req, timeout=30) as resp:
        return json.load(resp)


def get_leyning(date_str):
    """Return the Hebcal leyning item for a Shabbat date (triennial included)."""
    url = (
        "https://www.hebcal.com/leyning?cfg=json"
        f"&start={date_str}&end={date_str}&triennial=on"
    )
    data = fetch_json(url)
    items = data.get("items", [])
    if not items:
        raise SystemExit(f"No leyning found for {date_str} — is it a Shabbat?")
    item = items[0]
    if not item.get("triennial"):
        raise SystemExit(
            f"{date_str} is '{item.get('name', {}).get('en')}' with no triennial "
            "reading (likely a holiday). Pass the parsha's regular Shabbat date, "
            "or extend this script to take explicit verse ranges."
        )
    return item


class _FootnoteStripper(HTMLParser):
    """Extract text, dropping <sup> markers and footnote elements entirely
    (footnotes contain nested tags, so regex stripping is unreliable)."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts = []
        self.skip_depth = 0

    def handle_starttag(self, tag, attrs):
        classes = dict(attrs).get("class", "")
        if self.skip_depth or tag == "sup" or "footnote" in classes:
            self.skip_depth += 1

    def handle_endtag(self, tag):
        if self.skip_depth:
            self.skip_depth -= 1

    def handle_data(self, data):
        if not self.skip_depth:
            self.parts.append(data)


def clean_english(text):
    """Strip Sefaria footnotes/markup from a JPS 2006 verse."""
    parser = _FootnoteStripper()
    parser.feed(text)
    text = "".join(parser.parts)
    text = re.sub(r"\s+", " ", text).strip()
    return re.sub(r"\s*[—–]\s*$", "", text)


def clean_hebrew(text):
    """Strip markup and parsha-break markers ({פ}/{ס}) from a Hebrew verse."""
    text = re.sub(r"<[^>]+>", "", text)
    text = html.unescape(text)
    text = re.sub(r"\{[פס]\}", "", text)
    return re.sub(r"[ \t]+", " ", text).strip()


def flatten(segments):
    """Sefaria returns a flat list for one chapter, nested lists across chapters."""
    if segments and isinstance(segments[0], list):
        return [v for chapter in segments for v in chapter]
    return list(segments)


def fetch_aliyah_text(book, begin, end):
    """Fetch (hebrew, english) verse lists for e.g. book='Genesis', '14:1'-'14:9'."""
    ref = f"{book} {begin}-{end}"
    url = (
        "https://www.sefaria.org/api/texts/"
        + urllib.parse.quote(ref)
        + "?context=0&commentary=0&ven="
        + urllib.parse.quote(ENGLISH_VERSION)
    )
    data = fetch_json(url)
    hebrew = [clean_hebrew(v) for v in flatten(data["he"])]
    english = [clean_english(v) for v in flatten(data["text"])]
    if not hebrew or len(hebrew) != len(english):
        raise SystemExit(
            f"Verse mismatch for {ref}: {len(hebrew)} Hebrew vs {len(english)} English"
        )
    return hebrew, english


# --- docx layout (matches the April 2026 build_source_sheet.py format) ---

def new_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.15
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    for old in rpr.findall(qn("w:rFonts")):
        rpr.remove(old)
    rpr.append(parse_xml(
        '<w:rFonts %s w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
        'w:cs="Times New Roman" w:eastAsia="Times New Roman"/>' % nsdecls("w")
    ))
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    return doc


def make_run(p, text, size_pt, bold=False, rtl=False):
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    if bold:
        run.bold = True
        rPr.append(parse_xml("<w:bCs %s/>" % nsdecls("w")))
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rPr.append(parse_xml(
        '<w:rFonts %s w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
        'w:cs="Times New Roman" w:eastAsia="Times New Roman"/>' % nsdecls("w")
    ))
    for old in rPr.findall(qn("w:szCs")):
        rPr.remove(old)
    rPr.append(parse_xml('<w:szCs %s w:val="%d"/>' % (nsdecls("w"), size_pt * 2)))
    if rtl:
        rPr.append(parse_xml("<w:rtl %s/>" % nsdecls("w")))
    return run


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, text, 14)


def add_aliyah_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    make_run(p, text, 12, bold=True)


def add_hebrew(doc, text):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml('<w:bidi %s w:val="1"/>' % nsdecls("w")))
    p.paragraph_format.space_before = Pt(4)
    make_run(p, text, 16, rtl=True)


def add_english(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    make_run(p, text, 12)


def build_handout(date_str, num_aliyot, out_dir):
    item = get_leyning(date_str)
    parsha = item["name"]["en"]
    triennial = item["triennial"]
    print(f"{date_str}: Parashat {parsha} (triennial year {item.get('triYear', '?')})")

    doc = new_document()
    add_title(doc, f"Parashat {parsha}")

    for n in range(1, num_aliyot + 1):
        aliyah = triennial.get(str(n))
        if not aliyah:
            print(f"  warning: no aliyah {n} in triennial data — stopping at {n - 1}")
            break
        book, begin, end = aliyah["k"], aliyah["b"], aliyah["e"]
        hebrew, english = fetch_aliyah_text(book, begin, end)
        print(f"  {ORDINALS[n]} Aliyah: {book} {begin}-{end} ({len(hebrew)} verses)")
        add_aliyah_header(doc, f"{ORDINALS[n]} Aliyah")
        for he, en in zip(hebrew, english):
            add_hebrew(doc, he)
            add_english(doc, en)

    y, m, d = date_str.split("-")
    out_path = Path(out_dir) / f"{y}.{m}.{d} - {parsha}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"  saved: {out_path}")
    return out_path


def main():
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--date", action="append", required=True,
                    help="Service date YYYY-MM-DD (repeatable)")
    ap.add_argument("--aliyot", type=int, default=3,
                    help="Number of triennial aliyot to include (default 3)")
    ap.add_argument("--out", default=str(Path(__file__).parent / "handouts"),
                    help="Output directory (default ./handouts)")
    args = ap.parse_args()

    for date_str in args.date:
        build_handout(date_str, args.aliyot, args.out)


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    main()
